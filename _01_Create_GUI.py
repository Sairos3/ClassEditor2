import tkinter as tk
from tkinter import messagebox
import threading
import json
import os
import shutil
import comtypes.client
from docx import Document
from _17_Create_WordFile import extract_kw_numbers
THEME_FILE = "updated_themes.json"
SCHEDULE_DOC = "Weekly_Class_Schedules.docx"

def read_schedule_from_file(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def load_themes():
    if os.path.exists(THEME_FILE):
        with open(THEME_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_themes(themes):
    with open(THEME_FILE, "w", encoding="utf-8") as f:
        json.dump(themes, f, indent=4, ensure_ascii=False)

def extract_class_info_from_docx(doc_path):
    if not os.path.exists(doc_path):
        return [], "Document not found. Please generate or provide the document."
    doc = Document(doc_path)
    class_info = []
    weekday_map = {"Montag": "Montag", "Dienstag": "Dienstag", "Mittwoch": "Mittwoch", "Donnerstag": "Donnerstag", "Freitag": "Freitag"}
    current_day = None
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 7:
                day_text = cells[0].text.strip()
                if day_text in weekday_map:
                    current_day = weekday_map[day_text]
                    continue
                if current_day:
                    theme = cells[1].text.strip()
                    instructor = cells[4].text.strip() if len(cells) > 4 else "Unknown"
                    if theme and instructor:
                        class_info.append((current_day, theme, instructor))
    if not class_info:
        return [], "No class information found in the document."
    return class_info, "Class information successfully extracted from the document."

def update_themes_in_docx(doc_path, updated_themes):
    doc = Document(doc_path)
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 7:
                theme = cells[1].text.strip()
                if theme in updated_themes:
                    cells[1].text = updated_themes[theme]
    doc.save(doc_path)

class ThemeEditorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ClassEditor")
        self.root.geometry("600x400")

        self.input_schedule_file = "schedule.txt"
        self.output_doc = "Weekly_Class_Schedules.docx"

        self.updated_themes = load_themes()
        self.class_info = []
        self.create_widgets()
        self.load_class_info()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def on_close(self):
        cleanup_thread = threading.Thread(target=self.cleanup_files)
        cleanup_thread.daemon = True
        cleanup_thread.start()
        cleanup_thread.join()
        self.root.quit()
        self.root.destroy()

    def cleanup_files(self):
        try:
            if os.path.exists("schedule.txt"):
                os.remove("schedule.txt")
            if os.path.exists("days"):
                if os.path.isdir("days"):
                    shutil.rmtree("days")
                else:
                    os.remove("days")
        except Exception as e:
            print(f"Error during cleanup: {e}")

    def create_widgets(self):
        self.class_info_label = tk.Label(self.root, text="Available Class Information:")
        self.class_info_label.pack(pady=5)

        self.class_info_listbox = tk.Listbox(self.root, width=70, height=10)
        self.class_info_listbox.pack(pady=5)

        self.rename_label = tk.Label(self.root, text="Edit Selected Class:")
        self.rename_label.pack(pady=5)

        self.rename_entry = tk.Entry(self.root, width=50)
        self.rename_entry.pack(pady=5)

        self.rename_button = tk.Button(self.root, text="Update Selected Classes", command=self.rename_class)
        self.rename_button.pack(pady=5)

        self.update_all_button = tk.Button(self.root, text="Update All Classes", command=self.update_all_themes)
        self.update_all_button.pack(pady=5)

        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=5)

        self.save_button = tk.Button(button_frame, text="Save PDF Document", command=self.save_updated_document)
        self.save_button.pack(side=tk.LEFT, padx=5)

        self.signature_checkbox = tk.Checkbutton(button_frame, text="Digital(No Signature)", command=self.toggle_signature)
        self.signature_checkbox.pack(side=tk.LEFT)

        self.status_label = tk.Label(self.root, text="", fg="green")
        self.status_label.pack(pady=5)

        self.include_signature = True



    def load_class_info(self):
        class_info, message = extract_class_info_from_docx(SCHEDULE_DOC)
        if class_info:
            self.class_info = class_info
            self.class_info_listbox.delete(0, tk.END)
            for day, theme, instructor in class_info:
                self.class_info_listbox.insert(tk.END, f"{day}: {theme}")
        else:
            messagebox.showerror("Error", message)

    def rename_class(self):
        selected = self.class_info_listbox.curselection()
        if not selected:
            messagebox.showwarning("Warning", "Please select a class to rename.")
            return
        
        selected_index = selected[0]
        day, old_theme, instructor = self.class_info[selected_index]

        if ' / ' in old_theme:
            theme_prefix = old_theme.split(' / ')[0]
            theme_suffix = old_theme.split(' / ')[1]
        else:
            theme_prefix = old_theme
            theme_suffix = ""

        new_theme_prefix = self.rename_entry.get().strip()

        if new_theme_prefix and new_theme_prefix != theme_prefix:
            new_theme = f"{new_theme_prefix} / {theme_suffix}" if theme_suffix else new_theme_prefix

            self.updated_themes[old_theme] = new_theme

            for i, (current_day, current_theme, current_instructor) in enumerate(self.class_info):
                if current_theme.startswith(theme_prefix):
                    if ' / ' in current_theme:
                        new_theme_for_instructor = f"{new_theme_prefix} / {current_theme.split(' / ')[1]}"
                        self.updated_themes[current_theme] = new_theme_for_instructor
                    else:
                        self.updated_themes[current_theme] = new_theme_prefix
                    # Update the class_info list
                    self.class_info[i] = (current_day, self.updated_themes[current_theme], current_instructor)

            save_themes(self.updated_themes)

            self.class_info[selected_index] = (day, new_theme, instructor)

            self.refresh_class_info_list()
            self.status_label.config(text=f"Theme '{old_theme}' renamed to '{new_theme}'.", fg="green")

            update_themes_in_docx(SCHEDULE_DOC, self.updated_themes)

        
    def refresh_class_info_list(self):
        self.class_info_listbox.delete(0, tk.END)
        for day, theme, instructor in self.class_info:
            self.class_info_listbox.insert(tk.END, f"{day}: {theme}")

    def update_all_themes(self):
        all_updated_themes = []
        for day, theme, instructor in self.class_info:
            updated_theme = self.updated_themes.get(theme, theme)
            all_updated_themes.append((day, updated_theme, instructor))
        
        self.class_info_listbox.delete(0, tk.END)
        for day, updated_theme, instructor in all_updated_themes:
            self.class_info_listbox.insert(tk.END, f"{day}: {updated_theme}")
        
        save_themes(self.updated_themes)
        self.status_label.config(text="All themes have been updated.", fg="green")

        update_themes_in_docx(SCHEDULE_DOC, self.updated_themes)
    
    def toggle_signature(self):
        self.include_signature = not self.include_signature

        from _17_Create_WordFile import create_schedule_document

        # 1. Recreate document with or without signature
        create_schedule_document(
            input_file=self.input_schedule_file,
            output_file=self.output_doc,
            include_signature=self.include_signature
        )

        # 2. Re-apply edited themes
        self.updated_themes = load_themes()
        if self.updated_themes:
            update_themes_in_docx(self.output_doc, self.updated_themes)

        status = "removed" if not self.include_signature else "included"
        self.status_label.config(
            text=f"Signature {status} in Word document.",
            fg="green"
        )

    def save_updated_document(self):
        self.updated_themes = load_themes()
        valid_themes = {theme: new_theme for theme, new_theme in self.updated_themes.items() if new_theme != "Praxisunterricht"}
        if not valid_themes:
            messagebox.showwarning("Warning", "No valid themes to save.")
            return
        message = update_themes_in_docx(SCHEDULE_DOC, valid_themes)
        
        self.save_as_pdf(SCHEDULE_DOC, self.include_signature)
        save_themes(valid_themes)


    def save_as_pdf(self, doc_path, include_signature):
        try:
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(doc_path))
            schedule_text = read_schedule_from_file('schedule.txt')
            kw_numbers = extract_kw_numbers(schedule_text)
            kw_label = "_".join(kw_numbers) if kw_numbers else "KW_Unknown"
            pdf_path = os.path.splitext(doc_path)[0] + f"_{kw_label}.pdf"
            
            if not include_signature:
                if doc.Paragraphs.Count > 1:
                    doc.Paragraphs(doc.Paragraphs.Count).Range.Delete()
            
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
            word.Quit()
            self.status_label.config(text=f"Document saved as PDF: {pdf_path}", fg="green")
        except Exception as e:
            self.status_label.config(text=f"Failed to save as PDF: {e}", fg="red")

if __name__ == "__main__":
    root = tk.Tk()
    app = ThemeEditorApp(root)
    root.mainloop()
