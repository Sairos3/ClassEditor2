import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx.shared import Inches
import re
import win32com.client
import random
import winreg

def extract_kw_numbers(schedule_text):
    try:
        pattern = r'KW:\s*(\d+)'
        matches = re.findall(pattern, schedule_text)
        return matches
    except Exception as e:
        print(f"An error occurred: {e}")
        return []

def get_word_username():
    username = None

    try:
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        username = word_app.UserName
        word_app.Quit()
    except Exception:
        username = None

    if not username or username.startswith("User"):
        for version in ["16.0", "15.0", "14.0"]:
            try:
                key_path = fr"SOFTWARE\Microsoft\Office\{version}\Word\Options"
                with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path) as key:
                    reg_username, _ = winreg.QueryValueEx(key, "UserName")
                    if reg_username and not reg_username.startswith("User"):
                        username = reg_username
                        break
            except FileNotFoundError:
                continue

    if not username or username.startswith("User"):
        try:
            identity_key = r"SOFTWARE\Microsoft\IdentityCRL\UserExtendedProperties"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, identity_key) as parent_key:
                subkey_name = winreg.EnumKey(parent_key, 0)
                with winreg.OpenKey(parent_key, subkey_name) as subkey:
                    # Try DisplayName OR FriendlyName
                    for field in ["DisplayName", "FriendlyName"]:
                        try:
                            name, _ = winreg.QueryValueEx(subkey, field)
                            if name and not name.startswith("User"):
                                username = name
                                break
                        except FileNotFoundError:
                            continue
        except Exception:
            pass

    if not username or username.startswith("User"):
        try:
            key_path = r"SOFTWARE\Microsoft\Office\16.0\Common\Identity\Identities"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path) as identities:

                i = 0
                while True:
                    try:
                        subkey = winreg.EnumKey(identities, i)
                        with winreg.OpenKey(identities, subkey) as user_key:

                            for field in ["DisplayName", "FriendlyName", "AccountName"]:
                                try:
                                    value, _ = winreg.QueryValueEx(user_key, field)
                                    if value and not value.startswith("User"):
                                        username = value
                                        break
                                except FileNotFoundError:
                                    continue

                            if username and not username.startswith("User"):
                                break

                        i += 1

                    except OSError:
                        break

        except Exception:
            pass

    if not username or username.startswith("User"):
        try:
            username = os.getlogin()
        except Exception:
            username = "Unknown"

    return username

def read_schedule_from_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_dates_and_class(schedule_text):
    class_name = re.search(r"Klasse:\s*(.*)", schedule_text)
    class_name = class_name.group(1).strip() if class_name else "Unbekannte Klasse"
    class_name = re.sub(r'[^A-Za-z0-9\s\-]', '', class_name)
    dates = [datetime.strptime(date, "%d.%m.%Y") for date in re.findall(r'(\d{2}\.\d{2}\.\d{4})', schedule_text)]
    if dates:
        return class_name, min(dates), max(dates)
    return class_name, None, None

def parse_class_info(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    return [
        (
            line.split(' | ')[0].split(' / ')[0].replace('Class: ', '').strip(),
            line.split(' | ')[0].split(' / ')[1].strip() if len(line.split(' | ')[0].split(' / ')) > 1 else '',
            line.split(' | ')[1].replace('Duration: ', '').strip()
        )
        for line in lines if ' | ' in line
    ]

def set_text_to_calibri(doc, font_size=10):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)

def set_table_borders(table, border_size=4):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()

            borders = tcPr.find(qn('w:tcBorders'))
            if borders is not None:
                for border in list(borders):
                    borders.remove(border)

            borders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                border_element = OxmlElement(f'w:{border}')
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), str(border_size))
                borders.append(border_element)

            tcPr.append(borders)

            v_align = tcPr.find(qn('w:vAlign'))
            if v_align is None:
                v_align = OxmlElement('w:vAlign')
                tcPr.append(v_align)
            v_align.set(qn('w:val'), 'center')

def add_day_schedule(day_name, class_info, table):
    if not class_info:
        print(f"No data found for {day_name}. Skipping.")
        return

    day_row = table.add_row()
    day_row.cells[0].text = day_name
    day_row.cells[0].paragraphs[0].bold = True
    day_row.cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for i in range(1, len(day_row.cells)):
        day_row.cells[0].merge(day_row.cells[i])

    valid_class_names = [info[0] for info in class_info if info[0] not in ["Sonderveranstaltung", "Praxisunterricht"]]

    for class_name, instructor, duration in class_info:
        row = table.add_row().cells

        if class_name in ["Sonderveranstaltung", "Praxisunterricht"]:
            instructor = class_name
            class_name = ""

        if instructor == "Praxisunterricht" and valid_class_names:
            class_name = random.choice(valid_class_names)

        if instructor == "Sonderveranstaltung" and valid_class_names:
            class_name = "IM NOT A MAGICIAN ENTER NEW THEMA!!!"

        if class_name and instructor:
            combined_text = f"{class_name} / {instructor}"
        else:
            combined_text = class_name if class_name else instructor

        if len(row) > 5:
            row[1].text = combined_text
            row[1].merge(row[2])
            row[1].merge(row[3])
            row[1].merge(row[4])
            row[1].merge(row[5])

            row[6].text = duration
            for paragraph in row[6].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_table_borders(table)

def set_table_properties(table, class_column_index, fixed_width=1.0):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()

            if i != class_column_index:
                cell.width = Inches(fixed_width)
            
            cell.text = cell.text.strip()
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            tcPr.append(OxmlElement('w:noWrap'))

def add_signature_line(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = False
    p_format = paragraph.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_element = paragraph._element
    p_pr = p_element.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    top_border = OxmlElement("w:top")
    top_border.set(qn("w:val"), "single")
    top_border.set(qn("w:sz"), "6")
    top_border.set(qn("w:space"), "1")
    top_border.set(qn("w:color"), "000000")
    p_borders.append(top_border)
    p_pr.append(p_borders)

def create_schedule_document(input_file, folder_path='days', output_file='Weekly_Class_Schedules.docx', existing_file_path=None, include_signature=True):
    schedule_text = read_schedule_from_file(input_file)
    class_name, oldest_date, newest_date = extract_dates_and_class(schedule_text)
    full_name = get_word_username()

    if not full_name:
        print("Could not retrieve full name from Word.")

    doc = Document(existing_file_path) if existing_file_path else Document()

    table1 = doc.add_table(rows=3, cols=6)
    table1.style = 'Table Grid'
    table1.cell(0, 0).merge(table1.cell(0, 1)).text = 'Name der/des Auszubildenden:'
    table1.cell(0, 2).merge(table1.cell(0, 5)).text = full_name
    table1.cell(1, 0).text = 'Ausbildungsjahr:'
    table1.cell(1, 1).text = str(datetime.now().year)
    table1.cell(1, 2).text = 'Abteilung:'
    table1.cell(1, 3).merge(table1.cell(1, 5)).text = class_name
    table1.cell(2, 0).text = 'Ausbildungswoche vom:'
    table1.cell(2, 1).text = oldest_date.strftime("%d.%m.%Y") if oldest_date else ''
    table1.cell(2, 2).text = 'bis:'
    table1.cell(2, 3).text = newest_date.strftime("%d.%m.%Y") if newest_date else ''
    table1.cell(2, 4).text = 'Nr.:'
    table1.cell(2, 5).text = extract_kw_numbers(schedule_text)
    cell_paragraph = table1.cell(2, 4).paragraphs[0]
    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    schedule_table = doc.add_table(rows=1, cols=8)
    schedule_table.style = 'Table Grid'
    hdr_cells = schedule_table.rows[0].cells
    hdr_cells[1].text = ('Betriebliche Tätigkeiten, Unterweisungen bzw. überbetriebliche Unterweisungen, '
                         'betrieblicher Unterricht, sonstige Schulungen, Themen des Berufsschulunterrichts')
    hdr_cells[6].text = 'Stunden'
    hdr_cells[7].text = 'Lfd. Nummer: Bezug zum Ausbildungs-rahmenplan (optionale Angabe)'
    hdr_cells[1].merge(hdr_cells[2]).merge(hdr_cells[3]).merge(hdr_cells[4]).merge(hdr_cells[5])

    days_of_week = ["Montag", "Dienstag", "Mittwoch", "Donnerstag", "Freitag"]
    for i, day in enumerate(days_of_week, 1):
        file_path = os.path.join(folder_path, f"{i}_{day.lower()}_schedule.txt")
        if not os.path.exists(file_path):
            print(f"Warning: File {file_path} not found. Skipping {day}.")
            continue
        class_info = parse_class_info(file_path)
        add_day_schedule(day, class_info, schedule_table)

    set_text_to_calibri(doc, font_size=10)
    for paragraph in hdr_cells[7].paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(7)
    
    doc.add_paragraph()

    if include_signature:
        signature_1 = doc.add_paragraph()
        add_signature_line(signature_1, "Datum, Unterschrift Auszubildende/r                              Datum, Unterschrift\n                                                                                                Ausbildender oder Ausbilderin/Ausbilder")

    doc.save(output_file)
